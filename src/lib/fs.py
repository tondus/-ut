import sys 
import os
import docx
from docx import Document
from docx.shared import Inches
import pprint as pp
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy as copy
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
import src.template.tas.tas as tas
import src.template.legacy.legacy as leg
import pandas as pd
import src.template.hr.hr as hr
import src.template.bw.bw as bw
import src.template.rfc.rfc as rfc
import src.template.filetransfer.file_transfer as ft
import src.template.filetransfer_sftp.filetransfer_sftp as ft_sftp
from src.lib.template import *
 

def get_paragraph(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    return new_para

def fix_picture_not_display(paragraph ,pic):
    ## adhoc fix : unreasonable 
    run =  paragraph.add_run()
    run.add_picture(pic , width=Inches(6))
    run.clear()
   
def insert_table_after(paragraph, table,data,pic):
    paragraph1 = get_paragraph(paragraph)
    paragraph2 = get_paragraph(paragraph1)
    paragraph1._p.addnext(table._tbl)

    run = paragraph2.add_run()
    run.font.color.rgb = RGBColor(255,0,0)
    run.add_text("* ")
    run.bold = True
    run = paragraph2.add_run()
    run.add_text("Mandatory")

    fix_picture_not_display(paragraph1,pic)
    

def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def remove_empty_line(doc):
    index = 0
    start_mandatory_index = 0
    end_mandatory_index = 0
    for p in doc.paragraphs:
        if "Interface" in p.text:
            start_mandatory_index = index
        if "#################" in p.text:
            end_mandatory_index = index
        index += 1
    
    for i in range(start_mandatory_index,end_mandatory_index+1):
        p = doc.paragraphs[i]
        if not p.text:
            remove_paragraph(p)

def create_document(filename,data,config,ricefw) -> int:
    doc = Document(filename)
    tables = doc.tables
     
    table_interface = tables[0]
    mark_paragraph = None
    
    index = 0
    mandatory_index = 0
    for p in doc.paragraphs:
        if "Interface" in p.text:
            mark_paragraph = p
            mandatory_index = index + 1
        index += 1

    paragraph_mandatory = doc.paragraphs[mandatory_index]
    run = paragraph_mandatory.add_run()
    run.add_text("##############################################")
    
    table = None
    pic = None
    for i in reversed(range(0, len(data['tables']))) :
        d = data['tables'][i]
        print("template",get_template_name(data['tables'][i]))
        match get_template_name(data['tables'][i]):
            case "tas":
                pic,text= tas.get_pic(d)
                table = tas.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case "rfc":
                pic,text = rfc.get_pic(d,config)
                table =  rfc.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case "hr":
                pic,text = hr.get_pic(d,config)
                table =  hr.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case "bw":
                pic,text = bw.get_pic(d,config)
                table =  bw.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case "file_transfer":
                pic,text = ft.get_pic(d,config)
                table =  ft.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case "file_transfer_sftp":
                pic,text = ft_sftp.get_pic(d,config)
                table =  ft_sftp.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
            case _:
                pic,text = leg.get_pic(d,config)
                table =  leg.add_data_to_table(copy.deepcopy(table_interface),d,config,ricefw,pic,text)
                
        insert_table_after(mark_paragraph,table,d,pic)
    
    # remove_empty_line(doc)
    table_interface._element.getparent().remove(table_interface._element) #remove default table
    remove_paragraph(paragraph_mandatory)
    write_output_file(doc,data)