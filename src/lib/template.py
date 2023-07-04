import configparser
import codecs
import sys 
import os
import docx
from docx import Document
from docx.shared import Inches
import pprint as pp
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy as copy
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.dml.color import ColorFormat
from root import ROOT_DIR
from docx.shared import RGBColor


col_ricef = 'ricef'
col_ut_description = 'ut_description'
col_source_system = 'source_system'
col_target_system = 'target_system'
col_odata = 'odata'
col_proxy = 'proxy'
col_rfc = 'rfc'
col_other1= 'other1'
col_cpi = 'cpi'
col_pi = 'pi'
col_other2 = 'other2'
col_inbound = 'inbound'
col_outbound = 'outbound'
col_synchronous = 'synchronous'
col_asynchronus = 'asynchronus'

img_tas_outbound = 'tas_outbound.png'
img_tas_inbound = 'tas_inbound.png'

checkbox_checked = "\u2612"
checkbox_unchecked = "\u2610"
checkbox_items = [col_odata,col_proxy,col_rfc,col_other1,col_cpi,col_pi,col_other2,col_inbound,col_outbound,col_synchronous,col_asynchronus]
root_images = ROOT_DIR+'/src/template/images/'
img_file_transfer = 'file_transfer'

def get_config():
    config = configparser.ConfigParser()
    config.read_file(codecs.open(ROOT_DIR+"/config.ini", "r", "utf8"))
    return config

def set_checkbox(checked,table,index): 
    value = "1" if checked else "0"
    table._element.xpath("//w14:checked")[1].set(qn("w14:val"), value)
    check_box_text = table._element.xpath("//w:tc//w:p//w:sdt//w:sdtContent//w:r//w:t")[index]
    check_box_text.getparent().replace(check_box_text, checkbox(checked))

def checkbox(checked):
    elm = OxmlElement('w:t')
    elm.text = checkbox_checked if checked else checkbox_unchecked
    return elm

def set_other_from_data(data,table):
    # table._tbl.xml
    # paragraph._p.xml
    # print(dir(row[5])) 
    # Use dir() to reveal _tc XML element
    # print(row[5]._tc.xml)
    # if data[col_ricef] == 'ZPIFII059':
    #     print('ZPIFII059',data[col_other1])
    if not pd.isna(data[col_other1]) :
        index = 15
        other_text = table._element.xpath("//w:tc//w:p//w:r//w:t")[index]
        other_text.text = data[col_other1]

    if not pd.isna(data[col_other2]) :
        index = 34
        other_text = table._element.xpath("//w:tc//w:p//w:r//w:t")[index]
        other_text.text = data[col_other2]

def set_checkbox_from_data(data,table):
    if data[col_odata]:set_checkbox(True,table,checkbox_items.index(col_odata))
    if data[col_proxy]:set_checkbox(True,table,checkbox_items.index(col_proxy))
    if data[col_rfc]:set_checkbox(True,table,checkbox_items.index(col_rfc))
    if data[col_cpi]:set_checkbox(True,table,checkbox_items.index(col_cpi))
    if data[col_pi]:set_checkbox(True,table,checkbox_items.index(col_pi))
    if not pd.isna(data[col_other1]) :set_checkbox(True,table,checkbox_items.index(col_other1))
    if not pd.isna(data[col_other2]):set_checkbox(True,table,checkbox_items.index(col_other2))
    if data[col_inbound]:set_checkbox(True,table,checkbox_items.index(col_inbound))
    if data[col_outbound]:set_checkbox(True,table,checkbox_items.index(col_outbound))
    if data[col_synchronous]:set_checkbox(True,table,checkbox_items.index(col_synchronous))
    if data[col_asynchronus]:set_checkbox(True,table,checkbox_items.index(col_asynchronus))

def set_source_target_from_data(data,table):
    if data['inbound']:
        table.cell(1, 1).text = data['source_system']
        table.cell(2, 1).text = data['target_system']
    else:
        table.cell(1, 1).text = data['source_system']
        table.cell(2, 1).text = data['target_system']

def set_picture_from_data(paragraph,pic):
    run = paragraph.add_run()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run.add_picture(copy.deepcopy(pic), width=Inches(6))

def set_process_from_data(paragraph,pic):
    paragraph.paragraph_format.space_before = Pt(10)
    # paragraph_process.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run()
    run.add_text("Process:")
    run.bold = True

def set_text_from_data_oneline(paragraph,text):
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = paragraph.add_run()
    run.add_text(text)

def set_text_from_data_bullet(paragraph,text):
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = paragraph.add_run()
    if text:
        run.add_text(text)

def get_template_name(data):
    name = None
    if "tas" in data[col_source_system].lower() or "tas" in data[col_target_system].lower():
        name  = 'tas'
    elif "file" in str(data[col_other1]).lower():
        name = 'file_transfer'
    elif 'HR' in data[col_source_system].upper() or 'HR' in data[col_target_system].upper():
        name  = 'hr'
    elif 'BW' in data[col_source_system].upper() or 'BW' in data[col_target_system].upper():
        name = 'bw'
    else:
        name = 'default'
    return name

def write_output_file(doc,data):
    if not os.path.exists('output/fs'):
        os.makedirs('output/fs')
    desc = data['tables'][0][col_ut_description]
    desc = desc.replace('/',',')
    ricefw = data['tables'][0][col_ricef]
    doc.save('output/fs/'+ricefw+'_'+desc.strip()+'.docx')
 