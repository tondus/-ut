import sys 
import os
import docx
from docx import Document
from docx.shared import Inches
import pprint as pp
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy as copy
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.dml.color import ColorFormat
from docx.shared import RGBColor

col_ricef = 'ricef'
col_source_system = 'source_system'
col_target_system = 'target_system'
col_odata = 'odata'
col_proxy = 'proxy'
col_rfc = 'rfc'
col_other1= 'other1'
col_cpi = 'cpi'
col_pi = 'pi'
col_other2 = 'other2'
col_inbound = 'inbound'
col_outbound = 'outbound'
col_synchronous = 'synchronous'
col_asynchronus = 'asynchronus'
col_si = 'si'
checkbox_checked = "\u2612"
checkbox_unchecked = "\u2610"
checkbox_items = ['odata','proxy','rfc','other1','cpi','pi','other2','inbound','outbound','synchronous','asynchronous']

def getTable(table):
    data = []
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return data

def checkbox(checked):
    elm = OxmlElement('w:t')
    elm.text = checkbox_checked if checked else checkbox_unchecked
    return elm

def get_paragraph(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    return new_para

def fix_picture_not_display(paragraph ,data):
    ## adhoc fix : unreasonable 
    pic =  r"C:\Developer\fs\source\images\master.png"
    run =  paragraph.add_run()
    run.add_picture(pic , width=Inches(6)) 
    run.clear()

def insert_table_after(paragraph, table,data):
    paragraph1 = get_paragraph(paragraph)
    paragraph2 = get_paragraph(paragraph1)
    paragraph1._p.addnext(table._tbl)

    run = paragraph2.add_run()
    run.font.color.rgb = RGBColor(255,0,0)
    run.add_text("* ")
    run.bold = True
    run = paragraph2.add_run()
    run.add_text("Mandatory")

    fix_picture_not_display(paragraph1,data)
    
 

# def set_checkbox(checked,cell): 
#     value = "1" if checked else "0"
#     # check_value = cell._element.xpath("//w:tc//w:p//w:sdt//w:sdtPr//w14:checkbox//w14:checked")[0]
#     # check_value.set(qn("w14:val"), value)
#     # check_box_text = cell._element.xpath("//w:tc//w:p//w:sdt//w:sdtContent//w:r//w:t")[0]
#     check_box_text = cell._element.xpath("//w:tc//w:p//w:sdt//w:sdtContent//w:r//w:t")[0]
#     print("aaaaaaaaa",len(cell._element.xpath("//w:tc//w:p//w:sdt//w:sdtContent//w:r//w:t")))
#     check_box_text.getparent().replace(check_box_text, checkbox(checked))

def set_checkbox(checked,table,index): 
    value = "1" if checked else "0"
    table._element.xpath("//w14:checked")[1].set(qn("w14:val"), value)
    check_box_text = table._element.xpath("//w:tc//w:p//w:sdt//w:sdtContent//w:r//w:t")[index]
    check_box_text.getparent().replace(check_box_text, checkbox(checked))

def add_data_to_table(table , data ,config,ricefw):
    pic = None 
    text = None
    # print(table._element.xml)
    # print(dir(row[1]))
    # print(row_first[2]._element.xml)
    # with open("xxx.xml", "w", encoding="utf-8") as f:
    #     f.write(table._element.xml)
    
    # set_checkbox(True,table,checkbox_items.index('odata'))
    # set_checkbox(True,table,checkbox_items.index('proxy'))
    # set_checkbox(True,table,checkbox_items.index('rfc'))
    # set_checkbox(True,table,checkbox_items.index('other1'))
    # set_checkbox(True,table,checkbox_items.index('cpi'))
    # set_checkbox(True,table,checkbox_items.index('pi'))
    # set_checkbox(True,table,checkbox_items.index('other2'))
    # set_checkbox(True,table,checkbox_items.index('inbound'))
    # set_checkbox(True,table,checkbox_items.index('outbound'))
    # set_checkbox(True,table,checkbox_items.index('synchronous'))
    # set_checkbox(True,table,checkbox_items.index('asynchronous'))

    if data[col_odata]:set_checkbox(True,table,checkbox_items.index('odata'))
    if data[col_proxy]:set_checkbox(True,table,checkbox_items.index('proxy'))
    if data[col_rfc]:set_checkbox(True,table,checkbox_items.index('rfc'))
    if data[col_cpi]:set_checkbox(True,table,checkbox_items.index('cpi'))
    if data[col_pi]:set_checkbox(True,table,checkbox_items.index('pi'))
    if data[col_other1]:set_checkbox(True,table,checkbox_items.index('other1'))
    if data[col_inbound]:set_checkbox(True,table,checkbox_items.index('inbound'))
    if data[col_outbound]:set_checkbox(True,table,checkbox_items.index('outbound'))
    if data[col_synchronous]:set_checkbox(True,table,checkbox_items.index('synchronous'))
    if data[col_asynchronus]:set_checkbox(True,table,checkbox_items.index('asynchronous'))
    
    
    if data['inbound']:
        table.cell(1, 1).text = data['source_system']
        table.cell(2, 1).text = data['target_system']
    else:
        table.cell(1, 1).text = data['source_system']
        table.cell(2, 1).text = data['target_system']

    pic =  r"C:\Developer\fs\source\images\master.png"
    

    row_content = table.row_cells(6)
    paragraph1 = row_content[0].add_paragraph()
    # paragraph2 = row_content[0].add_paragraph()
    paragraph3 = row_content[0].add_paragraph()
    files_paragraph = []
    for i in data['filenames']:
        files_paragraph.append(row_content[0].add_paragraph())

    paragraph4 = row_content[0].add_paragraph()
    paragraph5 = row_content[0].add_paragraph()
    
    run = paragraph1.add_run()
    paragraph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph1.paragraph_format.space_after = Pt(10)
    run.add_picture(copy.deepcopy(pic), width=Inches(6))
    
    # paragraph2.paragraph_format.space_before = Pt(10)
    # paragraph2.paragraph_format.space_after = Pt(10)
    # run = paragraph2.add_run()
    # run.add_text("Process:")
    # run.bold = True
    
    paragraph3.paragraph_format.left_indent = Inches(0.25)
    # paragraph3.paragraph_format.space_after = Pt(10)
    paragraph3.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = paragraph3.add_run()
    run.add_text(config['default']['master_line1'])
    for idx,p in enumerate(files_paragraph):
         p.paragraph_format.left_indent = Inches(0.75)
        #  p.paragraph_format.space_after = Pt(10)
         run = p.add_run()
         run.add_text(config['default']['bullet'])
         run.add_text("  ")
         run.add_text(data['filenames'][idx])


    paragraph4.paragraph_format.left_indent = Inches(0.25)
    # paragraph4.paragraph_format.space_after = Pt(10)
    paragraph4.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = paragraph4.add_run()
    run.add_text(config['default']['master_line2'])

    paragraph5.paragraph_format.left_indent = Inches(0.25)
    paragraph5.paragraph_format.space_after = Pt(10)
    paragraph5.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY
    run = paragraph5.add_run()
    run.add_text(config['default']['master_line3'])

    # run.add_break()
    # print(paragraph._p.xml)
    return table
def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def remove_empty_line(doc):
    index = 0
    start_mandatory_index = 0
    end_mandatory_index = 0
    for p in doc.paragraphs:
        if "Interface" in p.text:
            start_mandatory_index = index
        if "#################" in p.text:
            end_mandatory_index = index
        index += 1
    
    for i in range(start_mandatory_index,end_mandatory_index+1):
        p = doc.paragraphs[i]
        if not p.text:
            remove_paragraph(p)

def create_document(filename,data,config,ricefw) -> int:
    doc = Document(filename)
    tables = doc.tables
    
    table_interface = tables[0]
    mark_paragraph = None
    
    index = 0
    mandatory_index = 0
    for p in doc.paragraphs:
        if "Interface" in p.text:
            mark_paragraph = p
            mandatory_index = index + 1
        index += 1
    
    paragraph_mandatory = doc.paragraphs[mandatory_index]
    run = paragraph_mandatory.add_run()
    run.add_text("##############################################")
 
    for i in reversed(range(0, len(data['tables']))) :
        table = add_data_to_table(copy.deepcopy(table_interface),data['tables'][i],config,ricefw)
        insert_table_after(mark_paragraph,table,data['tables'][i])
    # remove_empty_line(doc)
    

    table_interface._element.getparent().remove(table_interface._element) #remove default table
    remove_paragraph(paragraph_mandatory)
    if not os.path.exists('output/master'):
        os.makedirs('output/master')
    print(data['tables'][0]['desc'])
    desc = data['tables'][0]['desc']
    desc = desc.replace('/','-')
    doc.save('output/master/'+ricefw+'_'+desc+'.docx')
 